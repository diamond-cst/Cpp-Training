from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "资料"
ASSET_DIR = OUT_DIR / "开题报告_assets"
OUT_FILE = OUT_DIR / "21张牌魔术_开题报告.docx"
UML_FILE = ASSET_DIR / "class_diagram.png"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_east_asia_font(run, font_name):
    run.font.name = "Calibri"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    set_east_asia_font(run, "宋体")


def set_table_borders(table, color="B8C2CC", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths_cm):
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(width)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_east_asia_font(run, "黑体")
    return p


def add_para(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r2 = p.add_run(text[len(bold_prefix):])
        runs = (r1, r2)
    else:
        runs = (p.add_run(text),)
    for run in runs:
        set_east_asia_font(run, "宋体")
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_east_asia_font(run, "宋体")
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_east_asia_font(run, "宋体")
    return p


def load_font(size, bold=False):
    candidates = [
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc",
        "/Library/Fonts/Arial Unicode.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            try:
                return ImageFont.truetype(path, size=size, index=0)
            except Exception:
                pass
    return ImageFont.load_default()


def wrap(draw, text, font, width):
    lines = []
    current = ""
    for chunk in text.split(" "):
        trial = chunk if not current else current + " " + chunk
        bbox = draw.textbbox((0, 0), trial, font=font)
        if bbox[2] - bbox[0] <= width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = chunk
    if current:
        lines.append(current)
    return lines


def box(draw, xy, title, members, fill, outline="#355C7D"):
    x1, y1, x2, y2 = xy
    title_font = load_font(26)
    body_font = load_font(20)
    draw.rounded_rectangle(xy, radius=10, fill=fill, outline=outline, width=3)
    draw.rectangle((x1, y1, x2, y1 + 48), fill=outline)
    draw.text((x1 + 16, y1 + 10), title, font=title_font, fill="white")
    y = y1 + 62
    for member in members:
        for line in wrap(draw, member, body_font, x2 - x1 - 30):
            draw.text((x1 + 16, y), line, font=body_font, fill="#1A202C")
            y += 26
        y += 2


def arrow(draw, start, end, label=None):
    draw.line((start, end), fill="#334155", width=3)
    x1, y1 = start
    x2, y2 = end
    if x2 >= x1:
        head = [(x2, y2), (x2 - 14, y2 - 8), (x2 - 14, y2 + 8)]
    else:
        head = [(x2, y2), (x2 + 14, y2 - 8), (x2 + 14, y2 + 8)]
    draw.polygon(head, fill="#334155")
    if label:
        font = load_font(18)
        lx = (x1 + x2) / 2
        ly = (y1 + y2) / 2 - 24
        draw.rounded_rectangle((lx - 56, ly - 4, lx + 56, ly + 24), radius=5, fill="#FFFFFF", outline="#CBD5E1")
        draw.text((lx - 46, ly), label, font=font, fill="#334155")


def create_uml():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (1800, 1180), "#F8FAFC")
    draw = ImageDraw.Draw(img)
    title_font = load_font(36)
    draw.text((60, 36), "21张牌魔术系统类图（初步设计）", font=title_font, fill="#0B2545")

    box(draw, (70, 120, 450, 390), "Card", [
        "- suit: Suit",
        "- rank: Rank",
        "- numericValue: int",
        "+ toString()",
        "+ getValue()",
        "+ operator==, <, <<, >>"
    ], "#EAF2FA")
    box(draw, (520, 120, 910, 390), "Deck<T>", [
        "- cards: T*",
        "- capacity: int",
        "- size: int",
        "+ addCard() / removeCard()",
        "+ shuffle() / clear()",
        "+ operator+, -, *, []"
    ], "#EAF2FA")
    box(draw, (1070, 110, 1540, 430), "MagicTrick", [
        "<<abstract>>",
        "+ initialize()",
        "+ performRound()",
        "+ displayState()",
        "+ reveal()",
        "+ saveState() / loadState()"
    ], "#E8EEF5", "#1F4D78")
    box(draw, (1010, 520, 1420, 850), "TwentyOneCardTrick", [
        "- workingDeck: Deck<Card>",
        "- pile1/2/3: Deck<Card>",
        "- currentRound, score",
        "+ dealIntoPiles()",
        "+ reorganizePiles()",
        "+ reveal()"
    ], "#F4F6F9")
    box(draw, (1450, 520, 1760, 850), "ConfigurableCardTrick", [
        "- deckSize, pileSize",
        "- revealIndex",
        "+ configure()",
        "+ initializeCards()"
    ], "#F4F6F9")
    box(draw, (560, 520, 920, 790), "TwentySevenCardTrick", [
        "- workingDeck: Deck<Card>",
        "- reveal index: 13",
        "+ performRound()",
        "+ reveal()"
    ], "#F4F6F9")
    box(draw, (80, 520, 440, 790), "Leaderboard", [
        "- records: vector<PlayerRecord>",
        "+ addOrUpdateRecord()",
        "+ getTopRecords()",
        "+ save() / load()"
    ], "#FEF7E8", "#7A5A00")
    box(draw, (520, 860, 900, 1100), "ReplayManager", [
        "- lines: vector<string>",
        "+ start()",
        "+ recordRound()",
        "+ recordReveal()",
        "+ exportHtml()"
    ], "#FEF7E8", "#7A5A00")
    box(draw, (980, 900, 1350, 1100), "NetworkGame", [
        "+ runMagicianServer()",
        "+ runAudienceClient()",
        "- SocketHandle"
    ], "#FEF7E8", "#7A5A00")

    arrow(draw, (450, 255), (520, 255), "存储")
    arrow(draw, (1070, 270), (910, 270), "使用")
    arrow(draw, (1215, 430), (1215, 520), "继承")
    arrow(draw, (1070, 310), (740, 520), "继承")
    arrow(draw, (1400, 430), (1540, 520), "继承")
    arrow(draw, (1010, 680), (910, 255), "组合")
    arrow(draw, (1010, 760), (900, 960), "记录")
    arrow(draw, (1010, 615), (440, 655), "更新成绩")
    img.save(UML_FILE)


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color in [
        ("Title", 20, "0B2545"),
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)

    header = section.header.paragraphs[0]
    header.text = "高级语言程序设计实训开题报告"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 116, 139)
        set_east_asia_font(run, "宋体")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("21张牌魔术 | 初步设计")


def add_info_table(doc):
    table = doc.add_table(rows=4, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    rows = [
        ("姓名", "陈思涛", "学号", "202330300551"),
        ("开题时间", "2026.4.21", "班级", "23级软件工程3班"),
        ("任课教师", "郭芬", "设计题目", "21张牌魔术"),
        ("报告阶段", "初步设计", "计划周期", "2026.4.21 - 2026.6.20"),
    ]
    for row, values in zip(table.rows, rows):
        for idx, text in enumerate(values):
            cell = row.cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if idx % 2 == 0:
                set_cell_shading(cell, "E8EEF5")
                set_cell_text(cell, text, bold=True, color="0B2545")
            else:
                set_cell_text(cell, text)
    set_table_width(table, [2.4, 4.3, 2.4, 4.3])


def add_module_table(doc):
    add_heading(doc, "2.2 系统功能模块划分", 2)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["模块", "主要职责", "对应类/文件"]):
        set_cell_shading(cell, "F2F4F7")
        set_cell_text(cell, text, bold=True, color="0B2545")
    data = [
        ("交互控制模块", "提供菜单、玩家信息录入、模式选择、牌堆选择、暂停恢复和结果确认。", "main_console.cpp、main_enhanced.cpp"),
        ("魔术算法模块", "完成初始化、循环发牌、观众选择、选中牌堆置中、三轮后定位揭示牌。", "MagicTrick、TwentyOneCardTrick、TwentySevenCardTrick、ConfigurableCardTrick"),
        ("牌与牌堆模块", "抽象牌面数据，维护动态牌堆，支持洗牌、合并、移除、复制和下标访问。", "Card、Deck<T>"),
        ("异常处理模块", "统一处理非法牌面、越界访问、非法输入、文件读写和游戏状态错误。", "Exceptions.h"),
        ("数据持久化模块", "保存/加载玩家进度，维护排行榜，记录游戏过程并导出回放。", "saveState/loadState、Leaderboard、ReplayManager"),
        ("扩展模块", "预留图形界面、声音、网络双人对战和更多自定义魔术。", "gui/、NetworkGame、assets/"),
    ]
    for row_data in data:
        cells = table.add_row().cells
        for idx, text in enumerate(row_data):
            set_cell_text(cells[idx], text, bold=(idx == 0))
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_table_width(table, [3.0, 7.0, 5.0])


def add_class_table(doc):
    add_heading(doc, "2.4 主要类及成员设计", 2)
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    for cell, text in zip(table.rows[0].cells, ["类/结构", "角色", "关键数据", "关键函数"]):
        set_cell_shading(cell, "F2F4F7")
        set_cell_text(cell, text, bold=True, color="0B2545")
    data = [
        ("Card", "表示数字牌或标准扑克牌。", "suit、rank、numericValue", "toString、getValue、operator==、operator<、operator<<、operator>>"),
        ("Deck<T>", "模板牌堆，负责动态数组管理。", "T* cards、capacity、size", "addCard、removeCard、shuffle、operator+、operator-、operator*、operator[]"),
        ("MagicTrick", "魔术抽象基类，统一不同魔术的接口。", "无具体数据成员", "initialize、performRound、displayState、reveal、saveState、loadState"),
        ("TwentyOneCardTrick", "21张牌魔术的具体实现。", "workingDeck、pile1/2/3、currentRound、score、playerName", "dealIntoPiles、reorganizePiles、performRound、reveal"),
        ("TwentySevenCardTrick", "27张牌变体，展示多态扩展。", "workingDeck、pile1/2/3、currentRound", "initializeCards、performRound、reveal"),
        ("ConfigurableCardTrick", "可配置15/21/27张牌变体。", "deckSize、pileSize、revealIndex", "configure、initializeCards、reorganizePiles"),
        ("Leaderboard", "维护玩家得分榜。", "records、filename、MAX_RECORDS", "addOrUpdateRecord、sortRecords、save、load"),
        ("ReplayManager", "记录和导出魔术过程。", "directory、lines、textFilename、htmlFilename", "recordRound、recordReveal、save、exportHtml"),
    ]
    for row_data in data:
        cells = table.add_row().cells
        for idx, text in enumerate(row_data):
            set_cell_text(cells[idx], text, bold=(idx == 0))
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_table_width(table, [3.2, 4.1, 4.1, 5.0])


def add_plan_table(doc):
    add_heading(doc, "三、工作计划（2026.4.21 至 2026.6.20）", 1)
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    for cell, text in zip(table.rows[0].cells, ["阶段", "时间", "主要任务", "阶段成果"]):
        set_cell_shading(cell, "F2F4F7")
        set_cell_text(cell, text, bold=True, color="0B2545")
    data = [
        ("需求分析与开题", "4.21 - 4.28", "阅读任务书，明确基本要求、评分档位、扩展目标；完成开题报告和初步设计。", "需求分析、方案设计、类图设计。"),
        ("核心类实现", "4.29 - 5.10", "实现Card、Deck<T>、MagicTrick和21张牌魔术主流程；完成输入校验和异常处理。", "可运行的控制台基础版。"),
        ("功能完善", "5.11 - 5.24", "实现保存/加载、多玩家存档、计分、排行榜、颜色显示、动画效果和27张/可配置变体。", "增强版控制台程序和核心测试。"),
        ("扩展与优化", "5.25 - 6.7", "完善GUI、网络对战、回放导出等加分功能；优化代码结构和用户体验。", "扩展功能演示版本。"),
        ("测试与报告", "6.8 - 6.20", "进行边界测试、异常测试、流程演示录制，整理实训报告和答辩材料。", "最终代码、实训报告、演示视频。"),
    ]
    for row_data in data:
        cells = table.add_row().cells
        for idx, text in enumerate(row_data):
            set_cell_text(cells[idx], text, bold=(idx == 0))
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_table_width(table, [2.4, 3.0, 7.2, 4.0])


def build_doc():
    create_uml()
    doc = Document()
    style_document(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("华南理工大学软件学院2023级\n高级语言程序设计实训开题报告")
    run.bold = True
    run.font.size = Pt(20)
    set_east_asia_font(run, "黑体")

    add_info_table(doc)

    add_heading(doc, "一、设计题目及功能描述", 1)
    add_para(doc, "本课题设计并实现一个“21张牌魔术”模拟系统。该魔术属于经典数学魔术：从21张互不相同的牌中让观众记住一张，程序将牌从上到下循环发成三叠，每叠7张；观众指出目标牌所在牌叠后，程序把该牌叠放在另外两叠中间并重新合并。连续执行三轮后，观众记住的牌必定位于整副牌第11张，程序据此揭示结果。")
    add_para(doc, "系统除完成观众模式下的完整魔术流程外，还计划支持魔术师练习模式、数字牌和标准扑克牌两种牌面、保存/加载、多玩家存档、计分与排行榜、彩色控制台、发牌动画、回放导出、27张牌和可配置牌数等变体。附加扩展方向包括Qt图形界面、声音提示和网络双人对战。")

    add_heading(doc, "二、设计思路", 1)
    add_heading(doc, "2.1 方案设计", 2)
    add_para(doc, "本系统采用“交互层 + 魔术算法层 + 数据结构层 + 持久化/扩展层”的分层方案。交互层负责菜单、输入输出和模式切换；魔术算法层通过抽象基类 MagicTrick 统一不同魔术流程；数据结构层由 Card 与 Deck<T> 管理牌面和牌堆；持久化与扩展层负责存档、排行榜、回放、GUI和网络功能。")
    add_para(doc, "核心算法按三轮循环设计：")
    add_number(doc, "初始化牌堆，支持数字1-21或从标准扑克牌中抽取不同牌面。")
    add_number(doc, "按“第一叠、第二叠、第三叠”循环发牌，将当前牌堆拆为三个等长牌堆。")
    add_number(doc, "读取观众选择的牌叠编号，并进行合法性校验；非法输入抛出并捕获 InvalidInputException。")
    add_number(doc, "将观众选择的牌叠置于中间位置，按未选牌叠 + 选中牌叠 + 未选牌叠的顺序合并。")
    add_number(doc, "重复三轮后，根据固定揭示位置输出结果；21张牌揭示索引为10，27张牌揭示索引为13，可配置魔术使用 (deckSize - 1) / 2。")
    add_para(doc, "通过基类指针或智能指针持有 MagicTrick 对象，运行时可选择 TwentyOneCardTrick、TwentySevenCardTrick 或 ConfigurableCardTrick，从而体现继承、多态和可扩展设计。Deck<T> 使用动态数组实现深拷贝、移动语义和析构释放，满足实训对构造函数、复制构造函数、析构函数与内存管理的要求。")

    add_module_table(doc)

    add_heading(doc, "2.3 类图设计", 2)
    add_para(doc, "类图重点体现三层关系：Card 是最小数据实体，Deck<T> 组合多个 Card；MagicTrick 是抽象接口，TwentyOneCardTrick、TwentySevenCardTrick 和 ConfigurableCardTrick 继承它并组合牌堆；Leaderboard、ReplayManager、NetworkGame 属于外围服务类，分别负责成绩、回放和网络扩展。")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(UML_FILE), width=Inches(6.5))
    cap = doc.add_paragraph("图1  21张牌魔术系统类图（初步设计）")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(71, 85, 105)
        set_east_asia_font(run, "宋体")

    add_class_table(doc)

    add_heading(doc, "2.5 数据文件设计", 2)
    add_para(doc, "系统涉及三类数据文件，均采用简单、可读、便于调试的文本或二进制兼容格式，并在文件读写时捕获 std::ios_base::failure 或自定义 FileIOException。")
    add_bullet(doc, "存档文件：保存玩家名、当前轮数、当前牌堆顺序、得分、模式参数、是否隐藏牌面等信息，存放在 saves/ 目录，可按玩家名区分。")
    add_bullet(doc, "排行榜文件：leaderboard.dat 保存前10名玩家记录，字段包括姓名、总分、游戏次数、正确次数、连续正确次数和时间戳。")
    add_bullet(doc, "回放文件：replays/ 目录保存每轮三叠牌内容、观众选择、最终揭示结果，并可导出 HTML，方便实训报告和演示视频取材。")
    add_bullet(doc, "资源文件：assets/ 预留卡牌图片、声音素材等，GUI和音效扩展可直接复用。")

    add_heading(doc, "2.6 主要流程设计", 2)
    add_para(doc, "程序主流程为：启动程序 -> 读取/创建玩家 -> 选择观众模式或魔术师练习模式 -> 选择21张牌、27张牌或自定义魔术 -> 初始化并展示牌堆 -> 执行三轮发牌和选择 -> 揭示目标牌 -> 询问是否正确 -> 更新分数、排行榜和回放 -> 选择继续、保存或退出。")
    add_para(doc, "异常流程包括：牌堆为空或大小不正确时终止当前局并提示重新初始化；叠数输入不在1到3时要求重新输入；文件保存或读取失败时提示用户更换文件名或重新创建存档；越界访问牌堆时抛出 OutOfBoundsException 并阻止程序继续使用错误数据。")

    add_plan_table(doc)

    add_heading(doc, "四、预期成果", 1)
    add_para(doc, "最终计划提交完整的C++项目源码、可运行程序、实训报告和演示视频。代码层面将体现类、继承、多态、模板、运算符重载、异常处理和动态内存管理；功能层面完成基础魔术流程，并尽量实现保存/加载、排行榜、动画、变体魔术、GUI、网络和回放等扩展。")

    doc.save(OUT_FILE)


if __name__ == "__main__":
    build_doc()
    print(OUT_FILE)
